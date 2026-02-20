from docx import Document
from datetime import date
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from doc_base import format_currency, extenso_brl, configurar_estilo, salvar_docx
import os
from utils import abreviar_texto


def add_coat_of_arms(doc):
    """Adiciona o brasão e cabeçalho oficial no documento."""
    # Caminho para o brasão
    brasao_path = os.path.join(os.path.dirname(__file__), "assets", "brasao.jpg")
    
    # Adicionar brasão centralizado
    p_brasao = doc.add_paragraph()
    p_brasao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_brasao.add_run()
    run.add_picture(brasao_path, width=Cm(4.0), height=Cm(3.82))  # Ajustar tamanho conforme necessário


def remover_bordas_tabela(table):
    """Remove todas as bordas de uma tabela DOCX via XML."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def fixar_largura_tabela(table, total_width_cm=16.0, column_widths_cm=None):
    """Força layout fixo, largura preferencial e grid de colunas."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    total_dxa = int(total_width_cm * 566.929)
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    if column_widths_cm:
        existing_grid = tbl.find(qn('w:tblGrid'))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        tblGrid = OxmlElement('w:tblGrid')
        for col_cm in column_widths_cm:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(col_cm * 566.929)))
            tblGrid.append(gridCol)
        tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)


def gerar_decreto(dados):
    doc = Document()
    configurar_estilo(doc)
    
    # Configurar margens oficiais
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.75)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    
    # ---------------------------------------------------------
    # CABEÇALHO COM BRASÃO
    # ---------------------------------------------------------
    add_coat_of_arms(doc)
    
    # ---------------------------------------------------------
    # TÍTULO - DECRETO N.º
    # ---------------------------------------------------------
    meses = {1:'janeiro', 2:'fevereiro', 3:'março', 4:'abril', 5:'maio', 6:'junho',
             7:'julho', 8:'agosto', 9:'setembro', 10:'outubro', 11:'novembro', 12:'dezembro'}
    hoje = dados.get('data_doc', date.today())
    
    p = doc.add_paragraph(f"DECRETO N.º {dados['numero']}, DE {hoje.day} DE {meses[hoje.month].upper()} DE {hoje.year}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'
    
   # doc.add_paragraph()  # Espaço
    
    # Ementa com recuo de 9cm
    p = doc.add_paragraph(f"Dispõe sobre a autorização para abertura de Crédito Adicional {dados['tipo_lei']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.left_indent = Cm(9.0)
    p.runs[0].bold = False
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    
    # ---------------------------------------------------------
    # PREÂMBULO
    # ---------------------------------------------------------
    p = doc.add_paragraph(f"        O Prefeito Municipal de {dados['municipio']}, Estado de São Paulo, usando de suas atribuições legais,")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    #doc.add_paragraph()  # Espaço
    
    p = doc.add_paragraph("         DECRETA:")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    #doc.add_paragraph()  # Espaço

    # ---------------------------------------------------------
    # ARTIGO 1º - CRÉDITOS
    # ---------------------------------------------------------
    texto_art1 = (
        f"         Art.1º Fica o Executivo Municipal autorizado a abrir no Departamento de Finanças/ "
        f"Divisão de Controle Financeiro da Prefeitura, um Crédito Adicional {dados['tipo_lei']} "
        f"na importância de {format_currency(dados['total_credito'])} ({extenso_brl(dados['total_credito'])}), "
        f"para atender as seguintes dotações:"
    )
    
    p = doc.add_paragraph(texto_art1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)

    # ---------------------------------------------------------
    # TABELA DE DOTAÇÕES
    # ---------------------------------------------------------
    # Detectar se os itens vêm da planilha (têm 'ficha')
    tem_ficha = any(str(item.get('ficha', '')).strip().isdigit() for item in dados['itens_credito'])
    num_cols = 3 if tem_ficha else 2

    table = doc.add_table(rows=0, cols=num_cols)
    table.style = "Table Grid"
    remover_bordas_tabela(table)

    if tem_ficha:
        widths = [Cm(0.75), Cm(11.5), Cm(2.5)]
        fixar_largura_tabela(table, total_width_cm=16.0, column_widths_cm=[0.75, 11.5, 2.5])
    else:
        widths = [Cm(13.5), Cm(2.5)]
        fixar_largura_tabela(table, total_width_cm=16.0, column_widths_cm=[13.5, 2.5])

    # Preencher itens
    for item in dados['itens_credito']:
        row = table.add_row()
        cells = row.cells
        label_full = item.get('label_docx', item.get('label', ''))
        label_full = abreviar_texto(label_full)

        if tem_ficha:
            ficha_val = str(item.get('ficha', '')).strip()
            prefix = f"{ficha_val} - "
            descricao_val = label_full[len(prefix):] if label_full.startswith(prefix) else label_full
            cells[0].paragraphs[0].text = ficha_val
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[1].paragraphs[0].text = descricao_val
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[2].paragraphs[0].text = format_currency(item.get('valor', 0))
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            cells[0].paragraphs[0].text = label_full
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[1].paragraphs[0].text = format_currency(item.get('valor', 0))
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = 'Times New Roman'

    # Aplicar larguras
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Total — última linha da tabela
    row_total = table.add_row()
    if tem_ficha:
        row_total.cells[0].paragraphs[0].text = ""
        row_total.cells[1].paragraphs[0].text = "TOTAL"
        row_total.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_total.cells[2].paragraphs[0].text = format_currency(dados['total_credito'])
        row_total.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        row_total.cells[0].paragraphs[0].text = "TOTAL"
        row_total.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_total.cells[1].paragraphs[0].text = format_currency(dados['total_credito'])
        row_total.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for c in row_total.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(8)

    # Espaço após tabela de crédito
    p_esp = doc.add_paragraph()
    p_esp.paragraph_format.space_before = Pt(0)
    p_esp.paragraph_format.space_after = Pt(0)

    # ---------------------------------------------------------
    # ARTIGO 2º - FONTES
    # ---------------------------------------------------------
    art_num = 2
    
    # Se houver ANULAÇÃO
    if dados['itens_anulacao']:
        texto_art2 = (
            f"         Art.{art_num}º Para cobertura do crédito autorizado no artigo anterior serão anuladas as seguintes dotações:"
        )
        p = doc.add_paragraph(texto_art2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
        
        # Tabela de anulação
        tem_ficha_a = any(str(i.get('ficha', '')).strip().isdigit() for i in dados['itens_anulacao'])
        num_cols_a = 3 if tem_ficha_a else 2
        table_a = doc.add_table(rows=0, cols=num_cols_a)
        table_a.style = 'Table Grid'
        remover_bordas_tabela(table_a)
        if tem_ficha_a:
            widths_a = [Cm(0.75), Cm(11.5), Cm(2.5)]
            fixar_largura_tabela(table_a, total_width_cm=16.0, column_widths_cm=[0.75, 11.5, 2.5])
        else:
            widths_a = [Cm(13.5), Cm(2.5)]
            fixar_largura_tabela(table_a, total_width_cm=16.0, column_widths_cm=[13.5, 2.5])

        for item in dados['itens_anulacao']:
            row = table_a.add_row()
            cells = row.cells
            label_full_a = item.get('label_docx', item.get('label', ''))
            if tem_ficha_a:
                ficha_val_a = str(item.get('ficha', '')).strip()
                prefix_a = f"{ficha_val_a} - "
                desc_a = label_full_a[len(prefix_a):] if label_full_a.startswith(prefix_a) else label_full_a
                cells[0].paragraphs[0].text = ficha_val_a
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].text = desc_a
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[2].paragraphs[0].text = format_currency(item.get('valor', 0))
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cells[0].paragraphs[0].text = label_full_a
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].text = format_currency(item.get('valor', 0))
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for c in cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                        r.font.name = 'Times New Roman'

        for row in table_a.rows:
            for idx, width in enumerate(widths_a):
                row.cells[idx].width = width

        # Total da anulação — última linha da tabela
        total_anulacao = sum(i['valor'] for i in dados['itens_anulacao'])
        row_total_a = table_a.add_row()
        if tem_ficha_a:
            row_total_a.cells[0].paragraphs[0].text = ""
            row_total_a.cells[1].paragraphs[0].text = "TOTAL"
            row_total_a.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_total_a.cells[2].paragraphs[0].text = format_currency(total_anulacao)
            row_total_a.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            row_total_a.cells[0].paragraphs[0].text = "TOTAL"
            row_total_a.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_total_a.cells[1].paragraphs[0].text = format_currency(total_anulacao)
            row_total_a.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for c in row_total_a.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(8)

        # Espaço após tabela de anulação
        p_esp = doc.add_paragraph()
        p_esp.paragraph_format.space_before = Pt(0)
        p_esp.paragraph_format.space_after = Pt(0)
        art_num += 1
    
    # Se houver superávit
    elif dados['val_sup'] > 0:
        texto_sup = (
            f"         Art.{art_num}º As despesas decorrentes deste decreto serão suportadas com recursos provenientes do "
            f"superávit financeiro, apurado na Prefeitura Municipal, nos termos do inc. I, § 1º, do art. 43, "
            f"da Lei n.º 4.320, de 17 de março de 1964, constituído pela diferença positiva entre o ativo e o passivo financeiro, "
            f"apurado no Balanço Patrimonial do exercício anterior, na importância de {format_currency(dados['val_sup'])} "
            f"({extenso_brl(dados['val_sup'])})."
        )
        p = doc.add_paragraph(texto_sup)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
       # doc.add_paragraph()  # Espaço
        art_num += 1
    
    # Se houver excesso de arrecadação
    elif dados['val_exc'] > 0:
        texto_exc = (
            f"         Art.{art_num}º As despesas decorrentes deste decreto serão suportadas com recursos provenientes de "
            f"excesso de arrecadação, nos termos do inciso II, § 1º, do artigo 43, da Lei nº 4.320, de 17 de março de 1964"
        )
        
        if dados.get('origem_recursos'):
            texto_exc += f", oriundos de {dados['origem_recursos']}"
        
        texto_exc += f", no valor de {format_currency(dados['val_exc'])} ({extenso_brl(dados['val_exc'])})."
        
        p = doc.add_paragraph(texto_exc)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
       # doc.add_paragraph()  # Espaço
        art_num += 1

    # ---------------------------------------------------------
    # ARTIGO 3º - PPA/LDO
    # ---------------------------------------------------------
    # Extrair números das leis
    ldo_info = dados.get('ldo', '')
    ppa_info = dados.get('ppa', '')
    
    texto_art3 = (
        f"         Art.{art_num}º As alterações promovidas nos artigos 1º e 2º do presente decreto, passam a fazer parte "
        f"da LDO {ldo_info} e PPA {ppa_info} visando atender ao disposto nos artigos 165 e 168 da CF, "
        f"artigo 2º da Instrução nº 2, do Tribunal de Contas do Estado de São Paulo, da LC 101, de 04 de maio de 2.000 e, "
        f"finalmente, para atender ao Projeto Audesp do Tribunal de Contas do Estado de São Paulo."
    )
    p = doc.add_paragraph(texto_art3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    art_num += 1

    # ---------------------------------------------------------
    # ARTIGO 4º - VIGÊNCIA
    # ---------------------------------------------------------
    p = doc.add_paragraph(f"       Art.{art_num}º Este decreto entra em vigor na data de sua publicação.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    # ---------------------------------------------------------
    # DATA E LOCAL
    # ---------------------------------------------------------
    data_extenso = f"        {dados['municipio']}, {hoje.day} de {meses[hoje.month]} de {hoje.year}."
    
    p = doc.add_paragraph(data_extenso)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    # 3 linhas em branco
   # doc.add_paragraph("\n\n")
    
    # ---------------------------------------------------------
    # ASSINATURA PREFEITO
    # ---------------------------------------------------------
    p = doc.add_paragraph(dados['prefeito'].upper())
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    #doc.add_paragraph("\n")
    
    # ---------------------------------------------------------
    # REGISTRO E PUBLICAÇÃO + ASSINATURA DA SECRETÁRIA
    # ---------------------------------------------------------
    texto_registro = (
        f"         Registrado e publicado na Secretaria Geral da Prefeitura Municipal de {dados['municipio']}, "
        f"Estado de São Paulo, em {hoje.day} de {meses[hoje.month]} de {hoje.year}."
    )
    p = doc.add_paragraph(texto_registro)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    # 3 linhas em branco
    #doc.add_paragraph("\n\n")
    
    # Assinatura da secretária
    p = doc.add_paragraph(dados.get('secretaria', 'RITA DE CÁSSIA CÔRTES FERRAZ').upper())
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)

    return salvar_docx(doc)
