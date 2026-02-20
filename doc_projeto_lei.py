from docx import Document
from datetime import date
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from doc_base import (
    configurar_estilo, adicionar_titulo, adicionar_ementa,
    adicionar_data_assinatura, salvar_docx,
    format_currency, extenso_brl
)
import os
from utils import abreviar_texto


def add_coat_of_arms(doc):
    """Adiciona o brasão e cabeçalho oficial no documento."""
    # Caminho para o brasão
    brasao_path = os.path.join(os.path.dirname(__file__), "assets", "brasao.jpg")
    
    # Adicionar brasão centralizado
    p_brasao = doc.add_paragraph()
    p_brasao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_brasao.add_run()
    run.add_picture(brasao_path, width=Cm(4.5))  # Ajustar tamanho conforme necessário
    

def remover_bordas_tabela(table):
    """Remove todas as bordas de uma tabela DOCX via XML."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def fixar_largura_tabela(table, total_width_cm=16.0, column_widths_cm=None):
    """Força layout fixo, largura preferencial e grid de colunas."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    total_dxa = int(total_width_cm * 566.929)
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    if column_widths_cm:
        existing_grid = tbl.find(qn('w:tblGrid'))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        tblGrid = OxmlElement('w:tblGrid')
        for col_cm in column_widths_cm:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(col_cm * 566.929)))
            tblGrid.append(gridCol)
        tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)


def classify_expense_type(itens_credito):
    """
    Classifica as despesas em 'de capital', 'de custeio' ou ambas.
    Critério: 
      - Custeio: começa com '3' (Ex: 3.3.90.36)
      - Capital: começa com '4' (Ex: 4.4.90.52)
    """
    tem_custeio = False
    tem_capital = False
    
    for item in itens_credito:
        # Pega o código da dotação (ficha)
        codigo = str(item.get('ficha', '')).strip()
        
        # Verifica o primeiro dígito da parte da natureza da despesa
        # Formato: XX.XX.XX.XX.XXX.XXXX.C.G.M.E.D...
        # O 'C' (categoria) é o que define capital (4) ou custeio (3)
        partes = codigo.split('.')
        if len(partes) > 6:
            categoria = partes[6]  # Categoria econômica
            if categoria == '3':
                tem_custeio = True
            elif categoria == '4':
                tem_capital = True
            
    if tem_custeio and tem_capital:
        return "de capital e de custeio"
    elif tem_capital:
        return "de capital"
    elif tem_custeio:
        return "de custeio"
    else:
        return "de custeio"  # Default


def gerar_projeto_lei(dados):
    doc = Document()
    configurar_estilo(doc)
    
    # Configurar margens oficiais
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.75)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    
    # ---------------------------------------------------------
    # CABEÇALHO COM BRASÃO
    # ---------------------------------------------------------
    add_coat_of_arms(doc)
    
    # ---------------------------------------------------------
    # TÍTULO
    # ---------------------------------------------------------
    p = doc.add_paragraph("PROJETO DE LEI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'
    
    #doc.add_paragraph()  # Espaço
    
    # Ementa com recuo de 9cm
    p = doc.add_paragraph(f"Dispõe sobre a abertura de Crédito Adicional {dados['tipo_lei']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.left_indent = Cm(9.0)
    p.runs[0].bold = False
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    
    #doc.add_paragraph()  # Espaço

    # ---------------------------------------------------------
    # PREÂMBULO
    # ---------------------------------------------------------
    p = doc.add_paragraph(f"O Prefeito Municipal de {dados['municipio']}, Estado de São Paulo:")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(1.27)
    
    #doc.add_paragraph()
    
    p = doc.add_paragraph("Faço saber que a Câmara Municipal decreta e eu sanciono a seguinte Lei:")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(1.27)
    #doc.add_paragraph()

    # ---------------------------------------------------------
    # ARTIGO 1º
    # ---------------------------------------------------------
    tipo_despesa_str = classify_expense_type(dados['itens_credito'])
    
    texto_art1 = (
        f"Art. 1º Fica o Executivo Municipal autorizado a abrir no Departamento de Finanças, desta Prefeitura, "
        f"um Crédito Adicional {dados['tipo_lei']}, na importância de {format_currency(dados['total_credito'])} "
        f"({extenso_brl(dados['total_credito'])}), para atender a despesa {tipo_despesa_str} "
        "para a seguinte dotação:"
    )
    
    p = doc.add_paragraph(texto_art1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    #doc.add_paragraph()  # Espaço antes da tabela

    # ---------------------------------------------------------
    # TABELA DE DOTAÇÕES (formato: código - elemento - departamento)
    # ---------------------------------------------------------
    # Detectar se os itens vêm da planilha (têm 'ficha')
    tem_ficha = any(str(item.get('ficha', '')).strip().isdigit() for item in dados['itens_credito'])
    num_cols = 3 if tem_ficha else 2

    table = doc.add_table(rows=0, cols=num_cols)
    table.style = "Table Grid"
    remover_bordas_tabela(table)

    if tem_ficha:
        widths = [Cm(1.0), Cm(11.5), Cm(2.5)]
        fixar_largura_tabela(table, total_width_cm=16.0, column_widths_cm=[0.75, 11.5, 2.5])
    else:
        widths = [Cm(13.5), Cm(2.5)]
        fixar_largura_tabela(table, total_width_cm=16.0, column_widths_cm=[13.5, 2.5])

    # Preencher itens
    for item in dados['itens_credito']:
        row = table.add_row()
        cells = row.cells
        label_full = item.get('label_docx', item.get('label', ''))
        label_full = abreviar_texto(label_full)

        if tem_ficha:
            ficha_val = str(item.get('ficha', '')).strip()
            prefix = f"{ficha_val} - "
            descricao_val = label_full[len(prefix):] if label_full.startswith(prefix) else label_full
            cells[0].paragraphs[0].text = ficha_val
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[1].paragraphs[0].text = descricao_val
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[2].paragraphs[0].text = format_currency(item.get('valor', 0))
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            cells[0].paragraphs[0].text = label_full
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[1].paragraphs[0].text = format_currency(item.get('valor', 0))
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = 'Times New Roman'

    # Aplicar larguras
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Total — última linha da tabela
    row_total = table.add_row()
    if tem_ficha:
        row_total.cells[0].paragraphs[0].text = ""
        row_total.cells[1].paragraphs[0].text = "Total"
        row_total.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_total.cells[2].paragraphs[0].text = format_currency(dados['total_credito'])
        row_total.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        row_total.cells[0].paragraphs[0].text = "Total"
        row_total.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_total.cells[1].paragraphs[0].text = format_currency(dados['total_credito'])
        row_total.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for c in row_total.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(8)

    # Espaço após tabela de crédito
    p_esp = doc.add_paragraph()
    p_esp.paragraph_format.space_before = Pt(0)
    p_esp.paragraph_format.space_after = Pt(0)

    # ---------------------------------------------------------
    # ARTIGO 2º - FONTES (EXCESSO OU SUPERÁVIT)
    # ---------------------------------------------------------
    art_num = 2
    
    if dados['val_exc'] > 0:
        # Texto base
        texto_exc = (
            f"Art. {art_num}º As despesas decorrentes desta lei serão suportadas com recursos provenientes de "
            f"excesso de arrecadação, nos termos do inciso II, § 1º, do artigo 43, da Lei nº 4.320, de 17 de março de 1964"
        )
        
        # Adicionar origem se fornecida
        if dados.get('origem_recursos'):
            texto_exc += f", oriundos de {dados['origem_recursos']}"
        
        texto_exc += f", no valor de {format_currency(dados['val_exc'])} ({extenso_brl(dados['val_exc'])})."
        
        p = doc.add_paragraph(texto_exc)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
        art_num += 1
        
    if dados['val_sup'] > 0:
        conector = ", ainda," if dados['val_exc'] > 0 else ""
        texto_sup = (
            f"Art. {art_num}º As despesas decorrentes desta lei serão suportadas{conector} com recursos provenientes do "
            f"superávit financeiro, apurado na Prefeitura Municipal, nos termos do inc. I, § 1º, do art. 43, "
            f"da Lei n.º 4.320, de 17 de março de 1964, constituído pela diferença positiva entre o ativo e o passivo financeiro, "
            f"apurado no Balanço Patrimonial do exercício anterior, na importância de {format_currency(dados['val_sup'])} "
            f"({extenso_brl(dados['val_sup'])})."
        )
        p = doc.add_paragraph(texto_sup)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
        #doc.add_paragraph()
        art_num += 1

    # Se houver ANULAÇÃO
    if dados['itens_anulacao']:
        conector = ", ainda," if (dados['val_exc'] > 0 or dados['val_sup'] > 0) else ""
        total_anul = sum(i['valor'] for i in dados['itens_anulacao'])
        texto_anul = (
             f"Art. {art_num}º As despesas decorrentes desta lei serão suportadas{conector} com recursos provenientes de "
             f"anulação parcial ou total de dotações orçamentárias, nos termos do inciso III, § 1º, do artigo 43, "
             f"da Lei nº 4.320, de 17 de março de 1964, na importância de {format_currency(total_anul)} "
             f"({extenso_brl(total_anul)}), conforme abaixo:"
        )
        p = doc.add_paragraph(texto_anul)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
        
        # Tabela de anulação
        tem_ficha_a = any(str(i.get('ficha', '')).strip().isdigit() for i in dados['itens_anulacao'])
        num_cols_a = 3 if tem_ficha_a else 2
        table_a = doc.add_table(rows=0, cols=num_cols_a)
        table_a.style = 'Table Grid'
        remover_bordas_tabela(table_a)
        if tem_ficha_a:
            widths_a = [Cm(1.0), Cm(11.5), Cm(2.5)]
            fixar_largura_tabela(table_a, total_width_cm=16.0, column_widths_cm=[0.75 11.5, 2.5])
        else:
            widths_a = [Cm(13.5), Cm(2.5)]
            fixar_largura_tabela(table_a, total_width_cm=16.0, column_widths_cm=[13.5, 2.5])

        for item in dados['itens_anulacao']:
            row = table_a.add_row()
            cells = row.cells
            label_full_a = item.get('label_docx', item.get('label', ''))
            if tem_ficha_a:
                ficha_val_a = str(item.get('ficha', '')).strip()
                prefix_a = f"{ficha_val_a} - "
                desc_a = label_full_a[len(prefix_a):] if label_full_a.startswith(prefix_a) else label_full_a
                cells[0].paragraphs[0].text = ficha_val_a
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].text = desc_a
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[2].paragraphs[0].text = format_currency(item.get('valor', 0))
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cells[0].paragraphs[0].text = label_full_a
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].text = format_currency(item.get('valor', 0))
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for c in cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                        r.font.name = 'Times New Roman'

        for row in table_a.rows:
            for idx, width in enumerate(widths_a):
                row.cells[idx].width = width

        # Total da anulação — última linha da tabela
        row_total_a = table_a.add_row()
        if tem_ficha_a:
            row_total_a.cells[0].paragraphs[0].text = ""
            row_total_a.cells[1].paragraphs[0].text = "Total"
            row_total_a.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_total_a.cells[2].paragraphs[0].text = format_currency(total_anul)
            row_total_a.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            row_total_a.cells[0].paragraphs[0].text = "Total"
            row_total_a.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_total_a.cells[1].paragraphs[0].text = format_currency(total_anul)
            row_total_a.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for c in row_total_a.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(8)

        # Espaço após tabela de anulação
        p_esp = doc.add_paragraph()
        p_esp.paragraph_format.space_before = Pt(0)
        p_esp.paragraph_format.space_after = Pt(0)
        art_num += 1

    # ---------------------------------------------------------
    # ARTIGO PPA / LDO
    # ---------------------------------------------------------
    texto_ppa = (
        f"Art. {art_num}º Fica o Poder Executivo Municipal autorizado, ainda, a proceder a inclusão do projeto previsto nesta Lei, "
        f"no valor de {format_currency(dados['total_credito'])} ({extenso_brl(dados['total_credito'])}), "
        f"no Plano Plurianual - {dados['ppa']} e na Lei de Diretrizes Orçamentárias - {dados['ldo']}, "
        "em vigência neste exercício, para atender às alterações introduzidas pelo Sistema Audesp do Tribunal de Contas do Estado de São Paulo."
    )
    p = doc.add_paragraph(texto_ppa)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    #doc.add_paragraph()
    art_num += 1

    # ---------------------------------------------------------
    # ARTIGO VIGÊNCIA
    # ---------------------------------------------------------
    p = doc.add_paragraph(f"Art. {art_num}º Esta lei entra em vigor na data de sua publicação.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    #doc.add_paragraph()

    # ---------------------------------------------------------
    # DATA E ASSINATURA
    # ---------------------------------------------------------
    meses = {1:'janeiro', 2:'fevereiro', 3:'março', 4:'abril', 5:'maio', 6:'junho',
             7:'julho', 8:'agosto', 9:'setembro', 10:'outubro', 11:'novembro', 12:'dezembro'}
    hoje = dados.get('data_doc', date.today())
    data_extenso = f"Prefeitura Municipal de {dados['municipio']}, {hoje.day} de {meses[hoje.month]} de {hoje.year}."
    
    p = doc.add_paragraph(data_extenso)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    # 3 linhas em branco
    doc.add_paragraph("\n\n")
    
    p = doc.add_paragraph(dados['prefeito'].upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    
    p = doc.add_paragraph("Prefeito Municipal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)

    # ---------------------------------------------------------
    # JUSTIFICATIVA (Opcional, em nova página)
    # ---------------------------------------------------------
    if dados.get("justificativa"):
        doc.add_page_break()
        p = doc.add_paragraph("JUSTIFICATIVA")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].bold = True
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(14)
        doc.add_paragraph("\n")
        
        p = doc.add_paragraph(dados["justificativa"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)

    return salvar_docx(doc)
