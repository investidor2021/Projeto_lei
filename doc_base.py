from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import io
from datetime import date

# =============================
# FORMATAÇÕES E UTILITÁRIOS
# =============================

def format_currency(val):
    return f"R$ {val:,.2f}".replace(",", "v").replace(".", ",").replace("v", ".")

def extenso_brl(valor):
    inteiro = int(valor)
    centavos = int(round((valor - inteiro) * 100))
    txt = num2words(inteiro, lang='pt_BR') + (" real" if inteiro == 1 else " reais")
    if centavos > 0:
        txt += " e " + num2words(centavos, lang='pt_BR') + (" centavo" if centavos == 1 else " centavos")
    return txt

def configurar_estilo(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

def adicionar_titulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)

def adicionar_ementa(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(3.0)
    run = p.add_run(texto)
    run.italic = True
    run.font.size = Pt(10)

def adicionar_data_assinatura(doc, municipio, prefeito):
    doc.add_paragraph("\n")
    p = doc.add_paragraph(f"{municipio}, {date.today().strftime('%d de %B de %Y')}.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(prefeito.upper())
    run.bold = True
    doc.add_paragraph("Prefeito Municipal").alignment = WD_ALIGN_PARAGRAPH.CENTER

def salvar_docx(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================
# TABELAS DE CRÉDITO E ANULAÇÃO
# =============================

def adicionar_lista_creditos(doc, itens):
    for it in itens:
        p = doc.add_paragraph()
        run = p.add_run(f" {it['label_docx']} ................. {format_currency(it['valor'])}")
        run.font.size = Pt(9)

def adicionar_lista_anulacoes(doc, itens):
    for it in itens:
        p = doc.add_paragraph()
        run = p.add_run(f" Anulação da dotação {it['label_docx']} ................. {format_currency(it['valor'])}")
        run.font.size = Pt(9)
